from __future__ import annotations

import io
import json
import re
import tempfile
from pathlib import Path
from typing import Any

from api.utils.llm_client import LLMProvider, config as llm_config
from api.utils.logger import setup_logging
from api.utils.minio_conn import LexCompanionMinio
from deepagent.core.document_loaders.docdealing import (
    DocDealingLoader,
    DocParseResult,
    LayoutTextItem,
)
from deepagent.multiagent.legal_assistant.task_execution.contract_layout import (
    detect_placeholder_in_text,
    find_bbox_for_anchor,
)

logger = setup_logging()

_JSON_FENCE_RE = re.compile(r"^```(?:json)?\s*|\s*```$", re.MULTILINE)
_DOTS_RE = re.compile(r"(\.{3,}|_{3,}|…+)")
_EXTRACT_PROMPT = """Bạn phân tích mẫu hợp đồng/văn bản (markdown OCR) và liệt kê các ô/thông tin cần điền.

Trả về JSON:
{
  "fields": [
    {
      "id": "snake_case_id",
      "label": "Nhãn tiếng Việt",
      "required": true,
      "anchor_text": "đoạn text gần chỗ trống trong mẫu (copy nguyên nếu có)",
      "match_strategy": "placeholder_dots|explicit_placeholder|label_only"
    }
  ]
}

Quy tắc:
- id: snake_case, không dấu, duy nhất.
- Phát hiện chỗ trống: dấu chấm ..., gạch ___, [điền], {{...}}.
- Chỉ liệt kê field cần người dùng điền; bỏ tiêu đề cố định đã có sẵn nội dung đầy đủ.
- Tối đa 40 field; gộp field liên quan nếu quá nhiều.
- Tiếng Việt.
"""

_llm: LLMProvider | None = None

IMAGE_SUFFIXES = frozenset({".jpg", ".jpeg", ".png"})
DOCX_SUFFIX = ".docx"
PDF_SUFFIX = ".pdf"


def _get_llm() -> LLMProvider:
    global _llm
    if _llm is None:
        _llm = LLMProvider(llm_config)
    return _llm


def _parse_json(raw: str | None) -> dict[str, Any] | None:
    if not raw or not str(raw).strip():
        return None
    text = _JSON_FENCE_RE.sub("", str(raw).strip()).strip()
    try:
        data = json.loads(text)
        return data if isinstance(data, dict) else None
    except json.JSONDecodeError:
        return None


def load_document_bytes(*, tenant_id: str, location: str) -> bytes:
    minio = LexCompanionMinio()
    data = minio.get(tenant_id, location)
    if not data:
        raise FileNotFoundError(f"Object not found: {location}")
    return data


def parse_template_bytes(body: bytes, suffix: str) -> DocParseResult:
    return DocDealingLoader.parse_bytes_full(body, suffix=suffix)


def layout_items_to_dicts(items: list[LayoutTextItem]) -> list[dict[str, Any]]:
    return [
        {
            "text": it.text,
            "page_no": it.page_no,
            "bbox": dict(it.bbox),
            "coord_origin": it.coord_origin,
        }
        for it in items
    ]


def extract_form_fields(
    markdown: str,
    layout_items: list[LayoutTextItem] | None = None,
) -> list[dict[str, Any]]:
    raw = _get_llm().chat_text(
        [
            {
                "role": "user",
                "content": f"Nội dung mẫu:\n\n{(markdown or '')[:12000]}",
            }
        ],
        system_prompt=_EXTRACT_PROMPT,
        max_tokens=3000,
        temperature=0.1,
    )
    data = _parse_json(raw)
    fields = list((data or {}).get("fields") or [])
    normalized: list[dict[str, Any]] = []
    seen: set[str] = set()
    for f in fields:
        if not isinstance(f, dict):
            continue
        fid = re.sub(r"[^a-z0-9_]+", "_", str(f.get("id") or "field").lower()).strip("_")
        if not fid or fid in seen:
            fid = f"field_{len(seen) + 1}"
        seen.add(fid)
        anchor = str(f.get("anchor_text") or "").strip()
        field = {
            "id": fid,
            "label": str(f.get("label") or fid),
            "required": bool(f.get("required", True)),
            "anchor_text": anchor,
            "match_strategy": str(f.get("match_strategy") or "placeholder_dots"),
            "value": "",
        }
        if layout_items and anchor:
            bbox = find_bbox_for_anchor(anchor, layout_items)
            if bbox:
                field["bbox"] = bbox
        normalized.append(field)

    if not normalized and markdown:
        for m in _DOTS_RE.finditer(markdown):
            ctx_start = max(0, m.start() - 80)
            ctx = markdown[ctx_start : m.start()].strip().split("\n")[-1][:60]
            fid = f"blank_{len(normalized) + 1}"
            normalized.append(
                {
                    "id": fid,
                    "label": ctx or f"Ô trống {len(normalized) + 1}",
                    "required": True,
                    "anchor_text": markdown[ctx_start : m.end()],
                    "match_strategy": "placeholder_dots",
                    "value": "",
                }
            )
            if len(normalized) >= 15:
                break
    return normalized


def apply_values_to_docx(body: bytes, filled_values: dict[str, str], form_schema: list[dict]) -> bytes:
    from docx import Document

    doc = Document(io.BytesIO(body))
    id_to_field = {str(f.get("id") or ""): f for f in form_schema}

    for para in doc.paragraphs:
        text = para.text
        if not text:
            continue
        new_text = text
        for fid, value in filled_values.items():
            if not value:
                continue
            placeholder = "{{" + fid + "}}"
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, value)
            field = id_to_field.get(fid) or {}
            anchor = str(field.get("anchor_text") or "").strip()
            if anchor and anchor in new_text and detect_placeholder_in_text(new_text):
                new_text = _DOTS_RE.sub(value, new_text, count=1)
        if new_text != text:
            para.text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if not text:
                        continue
                    new_text = text
                    for fid, value in filled_values.items():
                        if not value:
                            continue
                        placeholder = "{{" + fid + "}}"
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, value)
                        field = id_to_field.get(fid) or {}
                        anchor = str(field.get("anchor_text") or "").strip()
                        if anchor and anchor in new_text and detect_placeholder_in_text(new_text):
                            new_text = _DOTS_RE.sub(value, new_text, count=1)
                    if new_text != text:
                        para.text = new_text

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def apply_values_to_image(
    body: bytes,
    filled_values: dict[str, str],
    form_schema: list[dict],
    *,
    page_size: tuple[int, int] | None = None,
) -> bytes:
    from PIL import Image, ImageDraw, ImageFont

    img = Image.open(io.BytesIO(body)).convert("RGB")
    draw = ImageDraw.Draw(img)
    w, h = img.size

    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
    except OSError:
        font = ImageFont.load_default()

    id_to_field = {str(f.get("id") or ""): f for f in form_schema}
    for fid, value in filled_values.items():
        if not value:
            continue
        field = id_to_field.get(fid) or {}
        bbox = field.get("bbox")
        if not isinstance(bbox, dict):
            continue
        l = float(bbox.get("l", 0))
        t = float(bbox.get("t", 0))
        r = float(bbox.get("r", w))
        b = float(bbox.get("b", h))
        # Docling coords may be in page space; scale if page_size given
        if page_size and page_size[0] > 0 and page_size[1] > 0:
            sx = w / page_size[0]
            sy = h / page_size[1]
            l, r = l * sx, r * sx
            t, b = t * sy, b * sy
        x0, y0 = int(l), int(t)
        x1, y1 = int(r), int(b)
        if x1 <= x0:
            x1 = min(w - 1, x0 + max(120, len(value) * 8))
        if y1 <= y0:
            y1 = min(h - 1, y0 + 20)
        draw.rectangle([x0, y0, x1, y1], fill=(255, 255, 255))
        draw.text((x0 + 2, y0 + 2), value[:200], fill=(0, 0, 0), font=font)

    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


def apply_values_to_pdf(body: bytes, filled_values: dict[str, str], form_schema: list[dict]) -> bytes:
    import fitz

    doc = fitz.open(stream=body, filetype="pdf")
    id_to_field = {str(f.get("id") or ""): f for f in form_schema}
    for fid, value in filled_values.items():
        if not value:
            continue
        field = id_to_field.get(fid) or {}
        bbox = field.get("bbox")
        if not isinstance(bbox, dict):
            continue
        page_no = max(1, int(bbox.get("page_no", 1))) - 1
        if page_no >= len(doc):
            page_no = 0
        page = doc[page_no]
        rect = fitz.Rect(
            float(bbox.get("l", 72)),
            float(bbox.get("t", 72)),
            float(bbox.get("r", 200)),
            float(bbox.get("b", 90)),
        )
        page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
        page.insert_text((rect.x0 + 2, rect.y0 + 12), value[:200], fontsize=10)

    out = doc.tobytes()
    doc.close()
    return out


def _apply_filled_values_to_markdown(
    markdown: str,
    filled_values: dict[str, str],
    form_schema: list[dict],
) -> str:
    """Chèn giá trị đã thu thập vào bản markdown tham chiếu (không sửa file gốc)."""
    text = (markdown or "").strip()
    if not text:
        return text

    id_to_field = {str(f.get("id") or ""): f for f in form_schema}
    for fid, value in filled_values.items():
        if not value:
            continue
        placeholder = "{{" + fid + "}}"
        if placeholder in text:
            text = text.replace(placeholder, value)
        field = id_to_field.get(fid) or {}
        anchor = str(field.get("anchor_text") or "").strip()
        if anchor and anchor in text and detect_placeholder_in_text(anchor):
            # Thay nhóm .../___ ngay trong đoạn anchor nếu có
            filled_anchor = _DOTS_RE.sub(value, anchor, count=1)
            text = text.replace(anchor, filled_anchor, 1)
        elif anchor and anchor in text:
            text = text.replace(anchor, f"{anchor} {value}".strip(), 1)

    return text


def compose_filled_docx_from_reference(
    *,
    template_markdown: str,
    filled_values: dict[str, str],
    form_schema: list[dict],
    source_suffix: str | None = None,
) -> bytes:
    """
    Soạn văn bản DOCX mới từ nội dung tham chiếu (Docling markdown) + field đã điền.
    File upload (PDF/ảnh/DOCX) không được chỉnh sửa trực tiếp — chỉ dùng để trích xuất.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Văn bản đã soạn thảo", level=1)
    if source_suffix:
        doc.add_paragraph(
            f"(Tham chiếu từ mẫu {source_suffix}; bản giao cho người dùng là DOCX mới.)"
        )

    merged_md = _apply_filled_values_to_markdown(
        template_markdown, filled_values, form_schema
    )
    if merged_md:
        doc.add_heading("Nội dung", level=2)
        for line in merged_md.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=3)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=4)
            else:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("Không trích xuất được nội dung từ mẫu tham chiếu.")

    filled_ids = [k for k, v in filled_values.items() if v]
    if form_schema and filled_ids:
        doc.add_heading("Thông tin đã điền", level=2)
        for f in form_schema:
            fid = str(f.get("id") or "")
            label = str(f.get("label") or fid)
            val = filled_values.get(fid) or ""
            if val:
                doc.add_paragraph(f"{label}: {val}")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def render_filled_document(
    *,
    body: bytes,
    suffix: str,
    filled_values: dict[str, str],
    form_schema: list[dict],
    template_markdown: str | None = None,
) -> tuple[bytes, str]:
    """
    Luôn xuất DOCX soạn mới. ``body``/``suffix`` chỉ phản ánh mẫu tham chiếu (parse đã xong).
    """
    del body  # file gốc không patch — chỉ dùng markdown đã trích
    suf = (suffix or "").lower()
    md = (template_markdown or "").strip()
    if not md:
        logger.warning(
            "render_filled_document: empty template_markdown, fallback field list only"
        )
    docx_bytes = compose_filled_docx_from_reference(
        template_markdown=md,
        filled_values=filled_values,
        form_schema=form_schema,
        source_suffix=suf or None,
    )
    return docx_bytes, ".docx"


def export_filled_as_docx(
    filled_values: dict[str, str],
    form_schema: list[dict],
    *,
    template_markdown: str = "",
) -> bytes:
    """Alias giữ tương thích — luôn soạn DOCX từ tham chiếu."""
    return compose_filled_docx_from_reference(
        template_markdown=template_markdown,
        filled_values=filled_values,
        form_schema=form_schema,
    )


def enrich_schema_with_layout(
    form_schema: list[dict[str, Any]],
    layout_items: list[LayoutTextItem],
) -> list[dict[str, Any]]:
    enriched = []
    for f in form_schema:
        item = dict(f)
        if not item.get("bbox"):
            anchor = str(item.get("anchor_text") or "").strip()
            if anchor:
                bbox = find_bbox_for_anchor(anchor, layout_items)
                if bbox:
                    item["bbox"] = bbox
        enriched.append(item)
    return enriched


def save_draft_to_minio(
    *,
    tenant_id: str,
    kb_id: str,
    body: bytes,
    suffix: str,
    draft_key_prefix: str,
    version: int,
) -> str:
    minio = LexCompanionMinio()
    key = f"{draft_key_prefix}/filled_v{version}{suffix}"
    put = minio.put(tenant_id, key, body)
    if put is None:
        raise RuntimeError("MinIO put failed for draft")
    return key
