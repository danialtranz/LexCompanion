from __future__ import annotations

import io
import re
import unicodedata
from typing import Any

_BLANK_RE = re.compile(r"(\.{3,}|_{3,}|…+|\[\s*\.{0,3}\s*\]|\{\{[^}]+\}\})")
# Ô điền thực sự: không tách từng ký tự … trong cụm ngày (7)….tháng…. năm…
_FILLABLE_BLANK_RE = re.compile(r"(\.{3,}|_{3,}|…{2,}|…\.(?=\s))")
_DATE_INLINE_CLAUSE_RE = re.compile(
    r"kể từ ngày\s*\(\d{1,2}\)[.…\s]*tháng[.…\s]*năm[.…\s]*",
    re.IGNORECASE,
)
_FOOTNOTE_REF_RE = re.compile(r"\((\d{1,2})\)\s*\.?\s*$")
_EXPLANATION_HEADING = "giải thích"
_DATE_PHRASE_RE = re.compile(r"ngày\s*tháng\s*năm", re.IGNORECASE)
_DATE_SLOT_RE = re.compile(r"^(ngày|tháng|năm)(\s*\(\d{1,2}\))?$", re.IGNORECASE)


def mask_date_clauses(text: str) -> str:
    """Ẩn dấu chấm trong cụm 'kể từ ngày (N)….tháng…. năm…' — không coi là ô điền."""

    def _repl(m: re.Match[str]) -> str:
        return re.sub(r"[.…]", " ", m.group(0))

    return _DATE_INLINE_CLAUSE_RE.sub(_repl, text or "")


def find_fillable_blank_matches(text: str) -> list[re.Match[str]]:
    """Các cụm trống thực sự cần điền (đã bỏ qua inline ngày/tháng/năm)."""
    return list(_FILLABLE_BLANK_RE.finditer(mask_date_clauses(text)))


def detect_blank_placeholder(text: str) -> bool:
    return bool(_BLANK_RE.search(text or ""))


def _norm_heading(text: str) -> str:
    t = unicodedata.normalize("NFC", (text or "").strip().lower())
    return re.sub(r"\s+", " ", t)


def is_explanation_heading(text: str) -> bool:
    return _norm_heading(text) == _EXPLANATION_HEADING


def is_footnote_explanation_block(text: str, *, in_explanation_section: bool) -> bool:
    """Đoạn chú thích (1), (2)... ở cuối mẫu — không phải ô điền."""
    t = (text or "").strip()
    if not t:
        return False
    if not in_explanation_section:
        return False
    if re.match(r"^\(\d{1,2}\)\s", t):
        return True
    if t.startswith("Ví dụ:") or t.startswith("Trường hợp "):
        return True
    if "ví dụ:" in t.lower() and re.search(r"\(\d{1,2}\)", t):
        return True
    return False


def _is_title_with_footnote(text: str) -> bool:
    """Tiêu đề mẫu kiểu 'ĐƠN XIN NGHỈ VIỆC (1)' — (N) là số chú thích, không phải ô điền."""
    t = (text or "").strip()
    if not _FOOTNOTE_REF_RE.search(t):
        return False
    if re.match(r"^(ĐƠN|CỘNG|BIÊN|TỜ|THỎA THUẬN|HỢP ĐỒNG)\b", t, re.IGNORECASE):
        return True
    letters = [c for c in t if c.isalpha()]
    if letters and sum(1 for c in letters if c.isupper()) / len(letters) >= 0.75 and len(t) < 100:
        return True
    return False


def is_date_related_label(label: str) -> bool:
    """Nhãn ô ngày/tháng/năm — bỏ qua, không hỏi user."""
    t = (label or "").strip().lower()
    if not t:
        return False
    if _DATE_PHRASE_RE.search(t):
        return True
    if _DATE_SLOT_RE.match(t):
        return True
    if re.fullmatch(r"\.?\s*tháng\.?", t):
        return True
    if re.fullmatch(r"\.?\s*năm\.?", t):
        return True
    if "lý do" in t:
        return False
    if re.search(r"ngày\s*\(\d{1,2}\)\s*$", t):
        return True
    if re.search(r"kể từ ngày\s*\(\d{1,2}\)\s*$", t):
        return True
    return False


def is_date_related_block(text: str) -> bool:
    """Dòng chủ yếu là ngày/tháng/năm — bỏ qua toàn bộ block."""
    t = (text or "").strip().lower()
    if not t:
        return False
    if _DATE_PHRASE_RE.search(t):
        return True
    if all(k in t for k in ("ngày", "tháng", "năm")):
        other_labels = (
            "lý do",
            "chức vụ",
            "bộ phận",
            "tên là",
            "kính gửi",
            "điện thoại",
            "fax",
            "email",
        )
        if not any(k in t for k in other_labels):
            return True
    return False


def detect_footnote_fill_ref(text: str) -> str | None:
    """
    Mẫu VN thường để (3), (4) ở cuối dòng là chỗ điền (chú thích ở mục Giải thích).
    Ví dụ: 'Người khởi kiện: (3)' → '(3)'
    """
    t = (text or "").strip()
    if not t or detect_blank_placeholder(t):
        return None
    if _is_title_with_footnote(t):
        return None
    m = _FOOTNOTE_REF_RE.search(t)
    if not m:
        return None
    # Bỏ qua nếu cả dòng chỉ là số thứ tự danh mục
    if re.fullmatch(r"\d{1,2}", t):
        return None
    return f"({m.group(1)})"


def classify_fill_block(
    text: str,
    *,
    in_explanation_section: bool,
) -> dict[str, Any]:
    """Phân loại block có cần điền hay không và lý do."""
    t = (text or "").strip()
    if not t:
        return {"needs_fill": False, "fill_reason": None, "footnote_ref": None}

    if is_footnote_explanation_block(t, in_explanation_section=in_explanation_section):
        return {"needs_fill": False, "fill_reason": None, "footnote_ref": None}

    if is_date_related_block(t):
        return {"needs_fill": False, "fill_reason": None, "footnote_ref": None}

    has_dots = detect_blank_placeholder(t)
    footnote_ref = detect_footnote_fill_ref(t)

    if has_dots and not in_explanation_section:
        return {
            "needs_fill": True,
            "fill_reason": "blank_dots",
            "footnote_ref": footnote_ref,
        }

    if footnote_ref and not in_explanation_section:
        return {
            "needs_fill": True,
            "fill_reason": "footnote_ref",
            "footnote_ref": footnote_ref,
        }

    # Dòng chỉ có số thứ tự trong danh mục đính kèm (1, 2, ...)
    if in_explanation_section is False and re.fullmatch(r"\d{1,2}", t):
        return {
            "needs_fill": True,
            "fill_reason": "list_item",
            "footnote_ref": None,
        }

    return {"needs_fill": False, "fill_reason": None, "footnote_ref": None}


def extract_structured_blocks(body: bytes) -> list[dict[str, Any]]:
    """
    Trích xuất cấu trúc DOCX gốc (paragraph + table cell).
    Gắn metadata phân biệt ô điền vs chú thích cuối mẫu.
    """
    from docx import Document

    doc = Document(io.BytesIO(body))
    blocks: list[dict[str, Any]] = []

    explanation_start: int | None = None
    for i, para in enumerate(doc.paragraphs):
        if is_explanation_heading(para.text or ""):
            explanation_start = i
            break

    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue
        in_expl = explanation_start is not None and i >= explanation_start
        fill_meta = classify_fill_block(text, in_explanation_section=in_expl)
        blocks.append(
            {
                "block_id": f"p_{i}",
                "kind": "paragraph",
                "text": text,
                "paragraph_index": i,
                "table_index": None,
                "row": None,
                "col": None,
                "in_explanation_section": in_expl,
                "has_blank_placeholder": detect_blank_placeholder(text),
                "needs_fill": fill_meta["needs_fill"],
                "fill_reason": fill_meta["fill_reason"],
                "footnote_ref": fill_meta["footnote_ref"],
            }
        )

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = (para.text or "").strip()
                    if not text:
                        continue
                    fill_meta = classify_fill_block(text, in_explanation_section=False)
                    blocks.append(
                        {
                            "block_id": f"t_{ti}_r_{ri}_c_{ci}",
                            "kind": "table_cell",
                            "text": text,
                            "paragraph_index": None,
                            "table_index": ti,
                            "row": ri,
                            "col": ci,
                            "in_explanation_section": False,
                            "has_blank_placeholder": detect_blank_placeholder(text),
                            "needs_fill": fill_meta["needs_fill"],
                            "fill_reason": fill_meta["fill_reason"],
                            "footnote_ref": fill_meta["footnote_ref"],
                        }
                    )
    return blocks


def blocks_to_llm_context(blocks: list[dict[str, Any]], *, max_chars: int = 14000) -> str:
    """Serialize blocks cho prompt LLM — chỉ phần thân đơn (không gửi Giải thích)."""
    lines: list[str] = []
    used = 0
    for b in blocks:
        if b.get("in_explanation_section"):
            continue
        line = (
            f"[{b.get('block_id')}] ({b.get('kind')}) "
            f"fill={b.get('needs_fill')} reason={b.get('fill_reason')} "
            f"ref={b.get('footnote_ref')} | {b.get('text', '')[:500]}"
        )
        if used + len(line) > max_chars:
            break
        lines.append(line)
        used += len(line) + 1
    return "\n".join(lines)


def fillable_blocks(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [b for b in blocks if b.get("needs_fill")]


def excerpt_for_group(
    blocks: list[dict[str, Any]],
    field_keys: set[str],
    form_schema: list[dict[str, Any]],
) -> str:
    """Đoạn text đọc từ DOCX gốc cho nhóm HITL."""
    id_to_block = {str(f.get("id") or ""): f.get("location") for f in form_schema}
    parts: list[str] = []
    for key in field_keys:
        loc = id_to_block.get(key) or {}
        block_id = str(loc.get("block_id") or "").strip()
        for b in blocks:
            if b.get("block_id") == block_id:
                parts.append(b.get("text") or "")
                break
        else:
            for f in form_schema:
                if str(f.get("id") or "") == key:
                    anchor = str(f.get("anchor_text") or "").strip()
                    if anchor:
                        parts.append(anchor)
                    break
    return "\n\n".join(p for p in parts if p.strip())[:4000]
