from __future__ import annotations

import io
import re
from typing import Any

from deepagent.multiagent.legal_assistant.task_execution.contract_layout import (
    detect_placeholder_in_text,
)
from deepagent.multiagent.legal_assistant.task_execution.docx_structure import (
    find_fillable_blank_matches,
)


def anchor_before_placeholder(text: str) -> str:
    """Anchor = phần nhãn trước cụm placeholder đầu tiên (không gồm dấu chấm/gạch)."""
    matches = find_fillable_blank_matches(text or "")
    if matches:
        return (text[: matches[0].start()]).rstrip()
    return (text or "").strip()


def _normalize_anchor(anchor: str, text: str) -> str:
    """Bỏ placeholder khỏi anchor nếu heuristic/LLM copy cả dòng."""
    a = (anchor or "").strip()
    if not a:
        return anchor_before_placeholder(text)
    matches = find_fillable_blank_matches(a)
    if matches:
        a = a[: matches[0].start()].rstrip()
    if a and a in text:
        return a
    fallback = anchor_before_placeholder(text)
    return fallback or a


def _replace_nth_blank(text: str, value: str, *, slot_index: int) -> str:
    """Thay cụm chấm/gạch thứ slot_index (0-based) trong text."""
    matches = find_fillable_blank_matches(text)
    if not matches:
        return text
    idx = max(0, min(slot_index, len(matches) - 1))
    m = matches[idx]
    return text[: m.start()] + value + text[m.end():]


def _replace_first_fillable_blank(text: str, value: str) -> str:
    matches = find_fillable_blank_matches(text)
    if not matches:
        return text
    m = matches[0]
    return text[: m.start()] + value + text[m.end():]


def _apply_strategy_to_text(
    text: str,
    *,
    value: str,
    field_id: str,
    anchor: str,
    strategy: str,
    slot_index: int | None = None,
) -> str:
    if not value or not text:
        return text

    strat = (strategy or "replace_trailing_blank").strip().lower()

    # ---- slot-aware: chiến lược blank đều nhận slot_index ----
    # replace_placeholder_dots_nth là chiến lược tường minh cho multi-blank
    if strat in ("replace_placeholder_dots_nth", "placeholder_dots_nth"):
        return _replace_nth_blank(text, value, slot_index=slot_index or 0)

    if strat == "replace_token":
        token = "{{" + field_id + "}}"
        if token in text:
            return text.replace(token, value)
        return text

    if strat == "replace_footnote_ref":
        m = re.search(r"\((\d{1,2})\)\s*\.?\s*$", text)
        if m:
            return text[: m.start()] + value
        # Re-apply trên bản đã patch (mất (N) ở cuối) — thay toàn bộ phần sau nhãn, không nối thêm.
        label_prefix = re.sub(r"\(\d{1,2}\)\s*\.?\s*$", "", (anchor or "")).rstrip()
        if not label_prefix:
            label_prefix = re.sub(r"\(\d{1,2}\)\s*\.?\s*$", "", text).rstrip()
        if label_prefix and text.startswith(label_prefix):
            return label_prefix + value
        return f"{text} {value}".strip() if text else value

    if strat == "replace_list_item":
        if re.fullmatch(r"\d{1,2}", text.strip()):
            return f"{text.strip()}. {value}"
        return value if not text.strip() else f"{text} {value}".strip()

    if strat in ("replace_placeholder_dots", "placeholder_dots"):
        if not detect_placeholder_in_text(text):
            return text
        if slot_index is not None:
            return _replace_nth_blank(text, value, slot_index=slot_index)
        return _replace_first_fillable_blank(text, value)

    if strat == "append_after_anchor":
        if anchor and anchor in text:
            rest = text.split(anchor, 1)[1]
            if not rest.strip() or detect_placeholder_in_text(rest):
                return text.replace(anchor, f"{anchor}{value}", 1)
            return f"{text} {value}".strip()
        return f"{text} {value}".strip() if text else value

    if strat in ("replace_trailing_blank", "label_only"):
        effective_anchor = _normalize_anchor(anchor, text)
        if effective_anchor and effective_anchor in text:
            idx = text.index(effective_anchor) + len(effective_anchor)
            head, tail = text[:idx], text[idx:]
            if detect_placeholder_in_text(tail):
                if slot_index is not None:
                    tail = _replace_nth_blank(tail, value, slot_index=slot_index)
                else:
                    tail = _replace_first_fillable_blank(tail, value)
                return head + tail
            if not tail.strip():
                return f"{head} {value}".rstrip()
            # Đã điền rồi (re-apply) — không nối thêm
            if tail.strip() == value.strip():
                return text
            return f"{head} {value}".rstrip()
        if detect_placeholder_in_text(text):
            if slot_index is not None:
                return _replace_nth_blank(text, value, slot_index=slot_index)
            return _replace_first_fillable_blank(text, value)
        return text

    if strat == "fill_table_cell":
        if detect_placeholder_in_text(text):
            if slot_index is not None:
                return _replace_nth_blank(text, value, slot_index=slot_index)
            return _replace_first_fillable_blank(text, value)
        if not text.strip():
            return value
        return f"{text} {value}".strip()

    # fallback
    if "{{" + field_id + "}}" in text:
        return text.replace("{{" + field_id + "}}", value)
    if anchor and anchor in text and detect_placeholder_in_text(text):
        if slot_index is not None:
            return _replace_nth_blank(text, value, slot_index=slot_index)
        return _replace_first_fillable_blank(text, value)
    return text


def _set_paragraph_text(para, new_text: str) -> None:
    """Ghi text giữ run đầu (tránh mất style hoàn toàn)."""
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _locate_paragraph(doc, location: dict[str, Any]):
    from docx.text.paragraph import Paragraph

    kind = str(location.get("kind") or "paragraph")
    if kind == "table_cell":
        ti = int(location.get("table_index") or 0)
        ri = int(location.get("row") or 0)
        ci = int(location.get("col") or 0)
        if ti >= len(doc.tables):
            return None
        table = doc.tables[ti]
        if ri >= len(table.rows):
            return None
        row = table.rows[ri]
        if ci >= len(row.cells):
            return None
        paras = row.cells[ci].paragraphs
        return paras[0] if paras else None

    pidx = location.get("paragraph_index")
    if pidx is not None and 0 <= int(pidx) < len(doc.paragraphs):
        return doc.paragraphs[int(pidx)]

    block_id = str(location.get("block_id") or "")
    if block_id.startswith("p_"):
        try:
            idx = int(block_id[2:])
            if 0 <= idx < len(doc.paragraphs):
                return doc.paragraphs[idx]
        except ValueError:
            pass
    return None


def apply_field_strategies(
    body: bytes,
    filled_values: dict[str, str],
    form_schema: list[dict[str, Any]],
) -> bytes:
    """
    Patch DOCX gốc theo strategy + location đã resolve.
    Không sửa markdown chunk.
    """
    from docx import Document

    if not body:
        return body

    doc = Document(io.BytesIO(body))
    id_to_field = {str(f.get("id") or ""): f for f in form_schema}

    def _apply_sort_key(fid: str) -> tuple[str, int, str]:
        field = id_to_field.get(fid) or {}
        loc = field.get("location") or {}
        block_id = str(loc.get("block_id") or "")
        raw_slot = field.get("slot_index")
        slot = int(raw_slot) if raw_slot is not None else 0
        return (block_id, slot, fid)

    for fid, value in sorted(filled_values.items(), key=lambda kv: _apply_sort_key(kv[0])):
        if not (value or "").strip():
            continue
        field = id_to_field.get(fid) or {}
        location = field.get("location") or {}
        if not location:
            anchor = str(field.get("anchor_text") or "").strip()
            strategy = str(field.get("strategy") or field.get("match_strategy") or "replace_trailing_blank")
            for para in doc.paragraphs:
                text = para.text or ""
                if anchor and anchor in text:
                    new_text = _apply_strategy_to_text(
                        text, value=value.strip(), field_id=fid, anchor=anchor, strategy=strategy
                    )
                    if new_text != text:
                        _set_paragraph_text(para, new_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text or ""
                            if anchor and anchor in text:
                                new_text = _apply_strategy_to_text(
                                    text, value=value.strip(), field_id=fid, anchor=anchor, strategy=strategy
                                )
                                if new_text != text:
                                    _set_paragraph_text(para, new_text)
            continue

        para = _locate_paragraph(doc, location)
        if para is None:
            continue
        text = para.text or ""
        anchor = str(field.get("anchor_text") or "").strip()
        strategy = str(field.get("strategy") or field.get("match_strategy") or "replace_trailing_blank")
        raw_slot = field.get("slot_index")
        slot_index = int(raw_slot) if raw_slot is not None else None
        new_text = _apply_strategy_to_text(
            text, value=value.strip(), field_id=fid, anchor=anchor, strategy=strategy,
            slot_index=slot_index,
        )
        if new_text != text:
            _set_paragraph_text(para, new_text)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def apply_values_to_docx(
    body: bytes,
    filled_values: dict[str, str],
    form_schema: list[dict],
) -> bytes:
    """Alias tương thích — delegate sang apply_field_strategies."""
    return apply_field_strategies(body, filled_values, form_schema)
