from __future__ import annotations

import io
import os
from pathlib import Path

from deepagent.multiagent.legal_assistant.task_execution.apply_field_strategies import (
    apply_field_strategies,
)
from deepagent.multiagent.legal_assistant.task_execution.docx_field_extract import (
    extract_fields_from_docx_llm,
    extract_fields_heuristic,
)
from deepagent.multiagent.legal_assistant.task_execution.docx_structure import (
    detect_blank_placeholder,
    extract_structured_blocks,
    fillable_blocks,
    is_date_related_block,
    is_date_related_label,
)
from deepagent.multiagent.legal_assistant.task_execution.document_chunks import (
    refine_assessment_for_chunk,
)
from deepagent.multiagent.legal_assistant.task_execution.draft_preview import (
    attach_draft_preview,
)
from deepagent.multiagent.legal_assistant.task_execution.hitl_groups import (
    map_fields_to_group_indices,
)


def _make_sample_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Kính gửi Tòa án nhân dân ...........")
    doc.add_paragraph("Người khởi kiện: ...........")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_p10_fill_does_not_corrupt_date_clause():
    """Lý do chỉ thay cụm chấm sau lý do(8) — không chèn vào ngày/tháng/năm."""
    docx_path = Path(__file__).resolve().parent / "don-xin-nghi-viec-1.docx"
    body = docx_path.read_bytes()
    blocks = extract_structured_blocks(body)
    schema, _ = extract_fields_heuristic(blocks)
    p10_field = next(f for f in schema if f["location"]["block_id"] == "p_10")
    orig = next(b["text"] for b in blocks if b["block_id"] == "p_10")

    out = apply_field_strategies(body, {p10_field["id"]: "Tôi mệt"}, schema)
    out_text = next(b["text"] for b in extract_structured_blocks(out) if b["block_id"] == "p_10")

    assert "(7)….tháng…. năm…" in out_text
    assert "Tôi mệt" in out_text
    assert "Tôi mệt.tháng" not in out_text
    assert out_text.endswith("Tôi mệt.") or out_text.endswith("Tôi mệt")


def test_skip_date_related_fields():
    assert is_date_related_block("Ngày tháng năm sinh: ………………………")
    assert is_date_related_block("……, ngày …… tháng …… năm……")
    assert not is_date_related_block(
        "Tôi xin phép được thôi việc kể từ ngày (7)….tháng…. năm… với lý do(8): …"
    )
    assert is_date_related_label("ngày (7)")
    assert is_date_related_label("tháng")
    assert not is_date_related_label("với lý do(8)")


def test_detect_blank_placeholder():
    assert detect_blank_placeholder("Kính gửi Tòa án ...........")
    assert not detect_blank_placeholder("Kính gửi Tòa án nhân dân quận 1")


def test_extract_structured_blocks():
    body = _make_sample_docx()
    blocks = extract_structured_blocks(body)
    assert len(blocks) >= 2
    assert blocks[0]["kind"] == "paragraph"
    assert blocks[0]["has_blank_placeholder"] is True


def test_apply_field_strategies_replace_trailing_blank():
    body = _make_sample_docx()
    schema = [
        {
            "id": "court_name",
            "anchor_text": "Kính gửi Tòa án nhân dân",
            "strategy": "replace_trailing_blank",
            "location": {
                "block_id": "p_0",
                "kind": "paragraph",
                "paragraph_index": 0,
            },
        }
    ]
    out = apply_field_strategies(body, {"court_name": "quận Ba Đình"}, schema)
    blocks = extract_structured_blocks(out)
    assert "quận Ba Đình" in blocks[0]["text"]
    assert "..." not in blocks[0]["text"]


def test_map_fields_to_group_indices():
    schema = [{"id": "a"}, {"id": "b"}, {"id": "c"}]
    groups = [
        {"group_id": 0, "field_keys": ["a", "b"]},
        {"group_id": 1, "field_keys": ["c"]},
    ]
    mapping = map_fields_to_group_indices(schema, groups)
    assert mapping["a"] == 0
    assert mapping["c"] == 1


def test_parse_real_docx_structure_blocks():
    """Parse file DOCX thật — kiểm tra fillable vs chú thích Giải thích."""
    docx_path = (
        Path(__file__).resolve().parent / "don-xin-nghi-viec-1.docx"
    )
    assert docx_path.exists(), f"Missing sample docx: {docx_path}"

    body = docx_path.read_bytes()
    blocks = extract_structured_blocks(body)
    assert blocks, "Không trích xuất được block nào từ DOCX"

    fillable = fillable_blocks(blocks)
    assert fillable, "Không phát hiện block cần điền"

    fill_ids = {b.get("block_id") for b in fillable}
    # Không lấy chú thích (1), (14) ở mục Giải thích
    assert "p_37" not in fill_ids
    assert "p_48" not in fill_ids
    # Ngày tháng năm (p_7, p_20) bỏ qua — không fill
    assert "p_7" not in fill_ids
    assert "p_20" not in fill_ids
    assert "p_8" in fill_ids
    assert "p_11" in fill_ids

    print(f"\n[diag] total_blocks={len(blocks)} fillable={len(fillable)}")
    for i, block in enumerate(fillable, start=1):
        print(
            f"[diag] #{i} {block.get('block_id')} reason={block.get('fill_reason')} "
            f"ref={block.get('footnote_ref')} | {str(block.get('text') or '')[:200]}"
        )


def test_parse_real_docx_heuristic_fields_unique_locations():
    """
    Heuristic slot_index cho đơn xin nghỉ việc:
    - block_id+slot_index là unique key (nhiều field cùng block ok nếu slot khác)
    - tiêu đề (1) không được coi là ô điền
    - dòng nhiều chỗ trống (p_8, p_10, p_20) tách đúng slot
    """
    docx_path = (
        Path(__file__).resolve().parent / "don-xin-nghi-viec-1.docx"
    )
    body = docx_path.read_bytes()
    blocks = extract_structured_blocks(body)
    schema, groups = extract_fields_heuristic(blocks)

    assert len(schema) >= 12, f"Thiếu field, chỉ có {len(schema)}"

    # (block_id, slot_index) phải unique
    loc_slot_keys = [
        (
            str((f.get("location") or {}).get("block_id") or ""),
            f.get("slot_index"),
        )
        for f in schema
    ]
    assert len(loc_slot_keys) == len(set(loc_slot_keys)), \
        "Có field bị map trùng (block_id, slot_index)"

    block_ids = {
        str((f.get("location") or {}).get("block_id") or "") for f in schema
    }
    assert "p_2" not in block_ids, "Tiêu đề ĐƠN XIN NGHỈ VIỆC (1) không phải ô điền"

    assert "p_7" not in block_ids, "Ngày tháng năm sinh không cần điền"
    assert "p_20" not in block_ids, "Dòng ký ngày tháng năm không cần điền"

    # p_8: Chức vụ (5) + Bộ phận (6) trên cùng dòng
    p8_fields = [
        f for f in schema
        if str((f.get("location") or {}).get("block_id")) == "p_8"
    ]
    assert len(p8_fields) == 2, f"p_8 phải có 2 slot, hiện có {len(p8_fields)}"
    p8_labels = [f.get("label", "").lower() for f in p8_fields]
    assert any("chức vụ" in lbl for lbl in p8_labels)
    assert any("bộ phận" in lbl for lbl in p8_labels)

    # p_10: chỉ lý do — bỏ qua ngày/tháng/năm
    p10_fields = [
        f for f in schema
        if str((f.get("location") or {}).get("block_id")) == "p_10"
    ]
    assert len(p10_fields) == 1, f"p_10 chỉ còn lý do, hiện có {len(p10_fields)}"
    p10_label = p10_fields[0].get("label", "").lower()
    assert "lý do" in p10_label or "lý do" in p10_fields[0].get("anchor_text", "").lower()

    print(f"\n[diag] heuristic_fields={len(schema)} groups={len(groups)}")
    for i, f in enumerate(schema, start=1):
        loc = f.get("location") or {}
        req = 'R' if f.get('required') else 'O'
        print(
            f"[diag-field] #{i:02d} [{req}] {loc.get('block_id'):8s} "
            f"slot={str(f.get('slot_index')):4s} {f.get('strategy'):35s} | {str(f.get('label') or '')[:60]}"
        )


def test_apply_multi_slot_fill():
    """Fill số điện thoại + fax vào đúng 2 slot trên cùng 1 dòng."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Số điện thoại: …………………(nếu có); số fax: ………………….(nếu có)")
    buf = io.BytesIO()
    doc.save(buf)
    body = buf.getvalue()

    schema = [
        {
            "id": "so_dien_thoai",
            "anchor_text": "Số điện thoại: …………………(nếu có); số fax: ………………….(nếu có)",
            "strategy": "replace_placeholder_dots_nth",
            "slot_index": 0,
            "location": {"block_id": "p_0", "kind": "paragraph", "paragraph_index": 0},
        },
        {
            "id": "so_fax",
            "anchor_text": "Số điện thoại: …………………(nếu có); số fax: ………………….(nếu có)",
            "strategy": "replace_placeholder_dots_nth",
            "slot_index": 1,
            "location": {"block_id": "p_0", "kind": "paragraph", "paragraph_index": 0},
        },
    ]

    # Apply điện thoại trước
    out1 = apply_field_strategies(body, {"so_dien_thoai": "0901234567"}, schema)
    blocks1 = extract_structured_blocks(out1)
    assert "0901234567" in blocks1[0]["text"], "slot 0 không điền được"
    assert "…………………" in blocks1[0]["text"], "slot 1 phải còn nguyên sau khi fill slot 0"

    # Apply fax tiếp vào bản đã fill slot 0
    out2 = apply_field_strategies(out1, {"so_fax": "02812345678"}, schema)
    blocks2 = extract_structured_blocks(out2)
    assert "0901234567" in blocks2[0]["text"], "slot 0 bị mất sau fill slot 1"
    assert "02812345678" in blocks2[0]["text"], "slot 1 không điền được"
    assert "………………" not in blocks2[0]["text"], "còn sót cụm chấm sau fill cả 2 slot"


def test_apply_real_docx_fills_at_correct_positions():
    """Điền toàn bộ field — giá trị thay placeholder, không nối thừa ở cuối dòng."""
    docx_path = Path(__file__).resolve().parent / "don-xin-nghi-viec-1.docx"
    body = docx_path.read_bytes()
    blocks = extract_structured_blocks(body)
    schema, _ = extract_fields_heuristic(blocks)

    filled = {f["id"]: f"VAL_{i}" for i, f in enumerate(schema, start=1)}
    out = apply_field_strategies(body, filled, schema)
    out_blocks = extract_structured_blocks(out)
    by_id = {b["block_id"]: b["text"] for b in out_blocks}

    # replace_trailing_blank: thay chấm, không còn cụm ... dư
    p6 = by_id["p_6"]
    assert p6.startswith("Tôi tên là: VAL_4")
    assert "...." not in p6

    # multi-slot: cả 2 slot trên p_8 đều được điền
    p8 = by_id["p_8"]
    p8_ids = [f["id"] for f in schema if f["location"]["block_id"] == "p_8"]
    for fid in p8_ids:
        val = filled[fid]
        assert val in p8, f"{val} không có trong p_8"
    assert "...." not in p8

    # p_10: chỉ lý do
    p10_field = next(f for f in schema if f["location"]["block_id"] == "p_10")
    assert filled[p10_field["id"]] in by_id["p_10"]


def test_multi_slot_fields_are_required_for_hitl():
    """Dòng nhiều ô trống (p_8) phải required; ngày/tháng/năm bỏ qua."""
    docx_path = Path(__file__).resolve().parent / "don-xin-nghi-viec-1.docx"
    body = docx_path.read_bytes()
    blocks = extract_structured_blocks(body)
    schema, groups = extract_fields_heuristic(blocks)
    field_map = {str(f["id"]): 0 for f in schema}

    p8 = [f for f in schema if f["location"]["block_id"] == "p_8"]
    p10 = [f for f in schema if f["location"]["block_id"] == "p_10"]
    p20 = [f for f in schema if f["location"]["block_id"] == "p_20"]

    assert len(p8) == 2
    assert len(p10) == 1
    assert len(p20) == 0
    assert all(f.get("required") for f in p8 + p10)

    assessment = refine_assessment_for_chunk(
        {"missing_field_ids": [], "proposed_values": {}},
        form_schema=schema,
        field_chunk_index=field_map,
        chunk_index=0,
        merged_values={},
    )
    missing = set(assessment.get("missing_field_ids") or [])
    for f in p8 + p10:
        assert f["id"] in missing, f"slot {f['label']} phải nằm trong missing_field_ids"


def test_reapply_does_not_duplicate_values():
    """Mỗi vòng HITL/assess gọi lại apply — không được nối thêm giá trị."""
    docx_path = Path(__file__).resolve().parent / "don-xin-nghi-viec-1.docx"
    body = docx_path.read_bytes()
    blocks = extract_structured_blocks(body)
    schema, _ = extract_fields_heuristic(blocks)
    p6_field = next(f for f in schema if f["location"]["block_id"] == "p_6")

    filled = {p6_field["id"]: "HƯng"}
    state = {
        "template_mode": "docx_native",
        "_template_bytes": body,
        "working_docx_bytes": body,
        "form_schema": schema,
        "filled_values": filled,
    }

    for _ in range(5):
        state = attach_draft_preview(state)

    out_blocks = extract_structured_blocks(state["working_docx_bytes"])
    p6_text = next(b["text"] for b in out_blocks if b["block_id"] == "p_6")
    assert p6_text.count("HƯng") == 1, f"Giá trị bị lặp: {p6_text[:120]}"


def test_footnote_ref_reapply_on_patched_docx():
    """Đơn kháng cáo: re-apply trên bản đã patch (mất _template_bytes) không lặp giá trị."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Kính gửi: Tòa án nhân dân (1)")
    doc.add_paragraph("Người kháng cáo: (2)")
    buf = io.BytesIO()
    doc.save(buf)
    body = buf.getvalue()

    blocks = extract_structured_blocks(body)
    schema, _ = extract_fields_heuristic(blocks)
    filled = {f["id"]: "Tỉnh" if "p_0" in f["location"]["block_id"] else "Hưng" for f in schema}

    patched = apply_field_strategies(body, filled, schema)
    current = patched
    for _ in range(12):
        current = apply_field_strategies(current, filled, schema)

    out_blocks = extract_structured_blocks(current)
    for f in schema:
        bid = f["location"]["block_id"]
        txt = next(b["text"] for b in out_blocks if b["block_id"] == bid)
        val = filled[f["id"]]
        assert txt.count(val) == 1, f"Field {f['id']} bị lặp: {txt[:120]}"


def test_parse_real_docx_llm_fields_diagnostic():
    """
    Test chẩn đoán mapping field->location cho DOCX thật.
    Chỉ chạy khi có LLM key để tránh fail CI/local environment.
    """
    has_llm_key = bool(
        os.getenv("OPENAI_API_KEY")
        or os.getenv("ANTHROPIC_API_KEY")
        or os.getenv("DEEPSEEK_API_KEY")
    )
    if not has_llm_key:
        print("\n[diag] skip llm extraction (no API key in env)")
        return

    docx_path = (
        Path(__file__).resolve().parent / "don-xin-nghi-viec-1.docx"
    )
    body = docx_path.read_bytes()

    form_schema, hitl_groups, blocks = extract_fields_from_docx_llm(body)
    assert blocks, "LLM path không có structured blocks"
    assert form_schema, "LLM path không extract được field nào"
    assert hitl_groups, "LLM path không tạo được hitl groups"

    with_location = [f for f in form_schema if isinstance(f.get("location"), dict)]
    assert with_location, "Field extract ra nhưng chưa resolve location"

    print(f"\n[diag] llm_fields={len(form_schema)} groups={len(hitl_groups)}")
    for i, f in enumerate(form_schema[:30], start=1):
        loc = f.get("location") or {}
        print(
            f"[diag-field] #{i} id={f.get('id')} strategy={f.get('strategy')} "
            f"loc={loc} anchor={str(f.get('anchor_text') or '')[:120]}"
        )
